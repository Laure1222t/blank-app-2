import streamlit as st
import re
import time
import requests
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from difflib import SequenceMatcher
import jieba  # 用于中文分词，提高匹配精度

# 加载环境变量
load_dotenv()

# 设置页面配置
st.set_page_config(
    page_title="条款式政策比对分析工具",
    page_icon="📜",
    layout="wide"
)

# 自定义CSS
st.markdown("""
<style>
    .stButton>button {
        margin-top: 1rem;
    }
    .analysis-box {
        border: 1px solid #e0e0e0;
        border-radius: 5px;
        padding: 1rem;
        margin-top: 1rem;
        background-color: #f9f9f9;
    }
    .matched-clause {
        border-left: 4px solid #28a745;
        padding: 0.75rem;
        margin: 1rem 0;
        background-color: #f8fff8;
    }
    .non-compliant {
        border-left: 4px solid #dc3545;
        padding: 0.75rem;
        margin: 1rem 0;
        background-color: #fff5f5;
    }
    .difference-section {
        border-left: 4px solid #ffc107;
        padding: 0.75rem;
        margin: 0.5rem 0;
        background-color: #fffcf2;
    }
    .summary-box {
        border: 1px solid #007bff;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
        background-color: #f0f7ff;
    }
    .parse-info {
        font-size: 0.9rem;
        color: #6c757d;
        margin-top: 0.5rem;
    }
    .clause-item {
        padding: 0.5rem;
        margin: 0.25rem 0;
        border-radius: 3px;
        background-color: #f0f2f6;
    }
    .highlight-conflict { background-color: #ffeeba; padding: 2px 4px; border-radius: 3px; }
    .clause-box { border-left: 4px solid #007bff; padding: 10px; margin: 10px 0; background-color: #f8f9fa; }
    .compliance-ok { border-left: 4px solid #28a745; }
    .compliance-warning { border-left: 4px solid #ffc107; }
    .compliance-conflict { border-left: 4px solid #dc3545; }
    .model-response { background-color: #f0f2f6; padding: 15px; border-radius: 5px; margin: 10px 0; }
</style>
""", unsafe_allow_html=True)

# 初始化会话状态
if 'target_clauses' not in st.session_state:
    st.session_state.target_clauses = {}  # {条款号: 内容}
if 'compare_files' not in st.session_state:
    st.session_state.compare_files = {}  # {文件名: {条款: {}, 分析结果: {匹配结果: {}, 总结: ""}}}
if 'current_file' not in st.session_state:
    st.session_state.current_file = None
if 'api_key' not in st.session_state:
    st.session_state.api_key = os.getenv("QWEN_API_KEY", "")
if 'parse_precision' not in st.session_state:
    st.session_state.parse_precision = "中等"  # 解析精度

# 页面标题
st.title("📜 条款式政策比对分析工具")
st.markdown("严格按照'一、二、三……'和'（一）（二）（三）……'格式分割条款")
st.markdown("---")

# 条款提取设置
st.sidebar.subheader("条款提取设置")

# 解析精度设置
st.session_state.parse_precision = st.sidebar.select_slider(
    "条款解析精度",
    options=["宽松", "中等", "严格"],
    value=st.session_state.parse_precision,
    help="宽松：提取更多可能的条款；严格：只提取明确符合格式的条款"
)

# API配置
with st.expander("🔑 API 配置", expanded=not st.session_state.api_key):
    st.session_state.api_key = st.text_input("请输入Qwen API密钥", value=st.session_state.api_key, type="password")
    model_option = st.selectbox(
        "选择Qwen模型",
        ["qwen-turbo", "qwen-plus", "qwen-max"],
        index=0  # 默认使用轻量版
    )
    st.caption("提示：可从阿里云DashScope平台获取API密钥")

# 辅助函数：判断文本是否可能为附件内容
def is_likely_attachment(text):
    """判断文本是否可能为附件内容，返回True表示可能是附件"""
    if not text:
        return False
    
    # 附件通常有以下特征：
    # 1. 包含附件标识关键词
    attachment_keywords = ['附件', '附录', '附表', '附图', '附件一', '附录一', '附件列表']
    for kw in attachment_keywords:
        if kw in text:
            # 检查关键词附近是否有关联表述
            if re.search(f'{kw}[：: ]?[^\n]{0,20}(如下|如下所示|内容如下|包括|包含)', text):
                return True
    
    # 2. 包含文件格式扩展名
    file_extensions = r'\.(pdf|doc|docx|xls|xlsx|ppt|pptx|jpg|png|gif|zip|rar|txt)'
    if re.search(file_extensions, text, re.IGNORECASE):
        return True
    
    # 3. 包含附件编号格式
    if re.search(r'附件\s*[0-9一二三四五六七八九十]+[:：.、)]', text):
        return True
        
    return False

# 文本提取函数，跳过附件内容
def extract_text_from_pdf(file):
    """从PDF提取文本，优化中文处理，跳过附件内容"""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        attachment_count = 0
        skip_mode = False  # 是否进入跳过模式
        
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            
            # 处理中文空格和换行问题
            page_text = page_text.replace("  ", "").replace("\n", "").replace("\r", "")
            
            # 检查是否包含附件标识
            if not skip_mode and is_likely_attachment(page_text):
                skip_mode = True
                attachment_count += 1
                continue  # 跳过当前页
            
            # 如果已进入跳过模式，检查是否需要退出
            if skip_mode:
                # 连续多页空白或低信息量可能表示附件结束
                if len(page_text) < 50:
                    skip_mode = False
                continue  # 跳过附件页
            
            text += page_text
        
        # 提示跳过了多少附件内容
        if attachment_count > 0:
            st.info(f"已跳过 {attachment_count} 处可能的附件内容")
            
        return text
    except Exception as e:
        st.error(f"提取文本失败: {str(e)}")
        return ""

def split_into_clauses(text):
    """
    严格按照以下两种格式分割条款：
    1. 一、二、三、……格式（中文数字+顿号）
    2. （一）（二）（三）……格式（括号+中文数字+括号）
    """
    # 首先识别并提取所有符合两种格式的条款标记
    # 格式1: 一、二、三、...
    pattern1 = r'([一二三四五六七八九十百千]+、)'
    # 格式2: （一）（二）（三）...
    pattern2 = r'(\([一二三四五六七八九十百千]+\))'
    
    # 合并所有可能的条款标记
    all_markers = re.findall(pattern1, text) + re.findall(pattern2, text)
    
    if not all_markers:
        return []
    
    # 使用标记分割文本
    clauses = []
    prev_pos = 0
    
    # 遍历所有标记并分割文本
    for marker in all_markers:
        # 找到当前标记的位置
        pos = text.find(marker, prev_pos)
        if pos == -1:
            continue
            
        # 如果是第一个标记，且前面有内容，我们忽略前面的内容
        if prev_pos == 0 and pos > 0:
            prev_pos = pos
        
        # 提取当前标记到下一个标记之间的内容
        next_pos = text.find(all_markers[all_markers.index(marker)+1], pos) if all_markers.index(marker)+1 < len(all_markers) else len(text)
        
        clause = text[pos:next_pos].strip()
        if clause:
            clauses.append(clause)
            
        prev_pos = next_pos
    
    return clauses

def chinese_text_similarity(text1, text2):
    """计算中文文本相似度，使用分词后匹配"""
    # 使用jieba进行中文分词
    words1 = list(jieba.cut(text1))
    words2 = list(jieba.cut(text2))
    
    # 计算分词后的相似度
    return SequenceMatcher(None, words1, words2).ratio()

# PDF解析函数 - 按特定格式分割条款
def parse_pdf_by_clauses(file, precision="中等"):
    """解析PDF文件并严格按照指定格式分割条款"""
    try:
        with st.spinner("正在解析文件并按指定格式分割条款..."):
            # 提取文本并跳过附件
            full_text = extract_text_from_pdf(file)
            total_pages = len(PdfReader(file).pages)  # 获取总页数
            
            # 文本预处理
            full_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', full_text)  # 移除控制字符
            full_text = re.sub(r'\s+', ' ', full_text).strip()  # 统一空白字符
            
            # 按指定格式分割条款
            clauses_list = split_into_clauses(full_text)
            
            # 为条款添加编号并过滤
            clauses = {}
            for clause in clauses_list:
                # 提取条款编号（只处理指定的两种格式）
                num_match = None
                
                # 尝试从条款文本中提取编号
                if re.match(r'\([一二三四五六七八九十百千]+\)', clause):
                    num_match = re.match(r'\(([一二三四五六七八九十百千]+)\)', clause)
                elif re.match(r'[一二三四五六七八九十百千]+、', clause):
                    num_match = re.match(r'([一二三四五六七八九十百千]+)、', clause)
                
                if num_match:
                    clause_num = num_match.group(1)
                    # 清理条款内容，移除编号部分
                    clause_content = re.sub(r'^\([一二三四五六七八九十百千]+\)\s*', '', clause)
                    clause_content = re.sub(r'^[一二三四五六七八九十百千]+、\s*', '', clause_content)
                else:
                    # 不应该走到这里，因为split_into_clauses已经过滤了格式
                    continue
                
                # 根据精度过滤条款
                if precision == "严格" and len(clause_content) > 50:
                    clauses[clause_num] = clause_content.strip()
                elif precision == "中等" and len(clause_content) > 30:
                    clauses[clause_num] = clause_content.strip()
                elif precision == "宽松" and len(clause_content) > 20:
                    clauses[clause_num] = clause_content.strip()
            
            st.success(f"共解析 {total_pages} 页，按指定格式成功提取 {len(clauses)} 条条款")
            return clauses
            
    except Exception as e:
        st.error(f"文件解析错误: {str(e)}")
        return {}

# 调用Qwen API进行条款比对分析
def call_qwen_api(prompt, api_key, model="qwen-turbo"):
    """调用Qwen API进行条款比对分析"""
    if not api_key:
        st.error("请先配置API密钥")
        return None
    
    try:
        with st.spinner("正在分析条款..."):
            url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            data = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 1000
            }
            
            response = requests.post(url, headers=headers, json=data, timeout=60)
            
            if response.status_code == 200:
                response_data = response.json()
                if "choices" in response_data and len(response_data["choices"]) > 0:
                    return response_data["choices"][0]["message"]["content"]
                else:
                    st.error(f"API返回格式异常: {response_data}")
                    return None
            else:
                st.error(f"API调用失败: 状态码 {response.status_code}, 响应: {response.text}")
                return None
                
    except requests.exceptions.Timeout:
        st.error("API请求超时，请重试")
        return None
    except Exception as e:
        st.error(f"API请求错误: {str(e)}")
        return None

# 合规性分析函数
def analyze_clause_matches(target_clauses, compare_clauses, api_key, model):
    """按条款匹配进行合规性分析"""
    if not target_clauses or not compare_clauses:
        st.warning("缺少条款内容，无法进行分析")
        return None, None, 0, 0
    
    # 找到所有匹配的条款（条款号相同）
    all_matched_clause_nums = [num for num in target_clauses if num in compare_clauses]
    total_matched = len(all_matched_clause_nums)
    
    if not all_matched_clause_nums:
        # 尝试基于内容相似度匹配
        st.info("未找到编号匹配的条款，尝试基于内容相似度匹配...")
        target_list = [(num, content) for num, content in target_clauses.items()]
        compare_list = [(num, content) for num, content in compare_clauses.items()]
        
        matched_pairs = []
        used_indices = set()
        
        for i, (t_num, t_content) in enumerate(target_list):
            best_match = None
            best_ratio = 0.3  # 中文匹配阈值
            best_j = -1
            
            for j, (c_num, c_content) in enumerate(compare_list):
                if j not in used_indices:
                    ratio = chinese_text_similarity(t_content, c_content)
                    if ratio > best_ratio:
                        best_ratio = ratio
                        best_match = (c_num, c_content)
                        best_j = j
            
            if best_match:
                matched_pairs.append((t_num, best_match[0], best_ratio))
                used_indices.add(best_j)
        
        if matched_pairs:
            all_matched_clause_nums = [(t_num, c_num) for t_num, c_num, _ in matched_pairs]
            total_matched = len(matched_pairs)
            st.info(f"基于内容相似度找到 {total_matched} 条可能匹配的条款")
        else:
            st.info("未找到匹配的条款")
            return {}, "未找到匹配的条款，无法进行合规性分析。", 0, total_matched
    
    # 分析每个匹配的条款，筛选合规的
    compliant_results = {}
    non_compliant_results = {}
    
    with st.spinner(f"正在分析 {total_matched} 条匹配条款，筛选合规条款..."):
        progress_bar = st.progress(0)
        for i, item in enumerate(all_matched_clause_nums):
            # 处理两种匹配方式的结果
            if isinstance(item, tuple):
                t_num, c_num = item  # 相似度匹配的结果
            else:
                t_num = c_num = item  # 编号匹配的结果
                
            target_content = target_clauses[t_num]
            compare_content = compare_clauses[c_num]
            
            # 生成条款比对提示
            prompt = f"""
            请仔细分析以下两个中文条款的合规性：
            
            目标条款（第{t_num}条）：
            {target_content[:300]}
            
            待比对条款（第{c_num}条）：
            {compare_content[:300]}
            
            分析要求：
            1. 首先明确判断待比对条款是否符合目标条款要求（用"合规"或"不合规"开头）
            2. 指出两者的主要差异点（如无差异则说明一致）
            3. 分析差异可能带来的影响
            4. 注意中文法律/合同条款中常用表述的细微差别
            5. 用简洁的中文（不超过300字）输出分析结果
            """
            
            # 调用API分析
            result = call_qwen_api(prompt, api_key, model)
            if result:
                # 判断是否合规
                if result.strip().startswith("合规"):
                    compliant_results[t_num] = {
                        "target_num": t_num,
                        "compare_num": c_num,
                        "target": target_content,
                        "compare": compare_content,
                        "analysis": result,
                        "compliant": True
                    }
                else:
                    non_compliant_results[t_num] = {
                        "target_num": t_num,
                        "compare_num": c_num,
                        "target": target_content,
                        "compare": compare_content,
                        "analysis": result,
                        "compliant": False
                    }
            
            # 更新进度条
            progress_bar.progress((i + 1) / len(all_matched_clause_nums))
        
        # 限制只保留前50条合规条款
        max_analyze = 50
        final_compliant = dict(list(compliant_results.items())[:max_analyze])
        
        # 显示分析数量信息
        st.info(f"""
        分析完成：
        - 总匹配条款数：{total_matched} 条
        - 合规条款数：{len(compliant_results)} 条
        - 本次分析展示前 {min(len(compliant_results), max_analyze)} 条合规条款
        """)
    
    # 生成总体总结
    summary_prompt = f"""
    以下是目标政策文件与待比对文件中合规条款的分析结果：
    {json.dumps(final_compliant, ensure_ascii=False, indent=2)}
    
    额外信息：
    - 总匹配条款数：{total_matched} 条
    - 合规条款数：{len(compliant_results)} 条
    
    请基于以上分析，用简洁的中文（不超过300字）总结：
    1. 总体合规性情况
    2. 主要差异点汇总
    3. 简要的合规建议
    """
    
    summary = call_qwen_api(summary_prompt, api_key, model) or "无法生成总结，请检查API配置。"
    
    return final_compliant, summary, len(compliant_results), total_matched

# 生成Word文档
def generate_word_document(matched_results, summary, target_filename, compare_filename, total_compliant, total_matched):
    """生成Word格式分析报告"""
    try:
        doc = Document()
        
        # 标题
        title = doc.add_heading("政策文件条款比对分析报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_paragraph(f"目标政策文件: {target_filename}")
        doc.add_paragraph(f"待比对文件: {compare_filename}")
        doc.add_paragraph(f"分析日期: {time.strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"总匹配条款数: {total_matched}")
        doc.add_paragraph(f"合规条款数: {total_compliant}")
        doc.add_paragraph(f"本次报告分析条款数: {len(matched_results)}")
        doc.add_paragraph("")
        
        # 总体总结
        doc.add_heading("一、总体总结", level=1)
        for para in re.split(r'\n+', summary):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # 合规条款详细分析
        doc.add_heading("二、合规条款详细分析", level=1)
        
        for clause_num, details in matched_results.items():
            doc.add_heading(f"目标条款第{details['target_num']}条 vs 待比对条款第{details['compare_num']}条", level=2)
            
            p = doc.add_paragraph("目标条款内容：")
            p.style = 'Heading 3'
            doc.add_paragraph(details["target"])
            
            p = doc.add_paragraph("待比对条款内容：")
            p.style = 'Heading 3'
            doc.add_paragraph(details["compare"])
            
            p = doc.add_paragraph("分析结果：")
            p.style = 'Heading 3'
            for para in re.split(r'\n+', details["analysis"]):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name
            
    except Exception as e:
        st.error(f"生成Word文档失败: {str(e)}")
        return None

# 主界面布局
col1, col2 = st.columns([1, 2], gap="large")

with col1:
    st.subheader("目标政策文件")
    st.caption("作为基准的政策文件，系统将按'一、二、三'和'（一）（二）（三）'格式提取条款")
    target_file = st.file_uploader("上传目标政策文件 (PDF)", type="pdf", key="target")
    
    if target_file:
        # 解析目标文件特定格式条款
        st.session_state.target_clauses = parse_pdf_by_clauses(
            target_file, 
            precision=st.session_state.parse_precision
        )
        
        with st.expander(f"查看提取的条款 (共 {len(st.session_state.target_clauses)} 条)"):
            for num, content in st.session_state.target_clauses.items():
                display_text = content[:150] + "..." if len(content) > 150 else content
                st.markdown(f'<div class="clause-item"><strong>第{num}条:</strong> {display_text}</div>', unsafe_allow_html=True)
    
    # 多文件上传区域
    st.subheader("待比对文件")
    st.caption("可上传多个文件，系统将按'一、二、三'和'（一）（二）（三）'格式提取条款")
    compare_files = st.file_uploader(
        "上传待比对文件 (PDF)", 
        type="pdf", 
        key="compare",
        accept_multiple_files=True
    )
    
    # 处理上传的多个文件
    if compare_files:
        for file in compare_files:
            if file.name not in st.session_state.compare_files:
                # 解析待比对文件特定格式条款
                clauses = parse_pdf_by_clauses(
                    file, 
                    precision=st.session_state.parse_precision
                )
                st.session_state.compare_files[file.name] = {
                    "clauses": clauses,
                    "matched_results": None,
                    "summary": None,
                    "total_compliant": 0,
                    "total_matched": 0
                }
                st.success(f"✅ 已添加 {file.name}，提取到 {len(clauses)} 条符合格式的条款")
    
    # 显示已上传的待比对文件列表
    if st.session_state.compare_files:
        st.subheader("已上传文件")
        for filename in st.session_state.compare_files.keys():
            col_a, col_b = st.columns([3, 1])
            with col_a:
                clause_count = len(st.session_state.compare_files[filename]["clauses"])
                if st.session_state.compare_files[filename]["total_compliant"] > 0:
                    st.markdown(f"- {filename} (条款数: {clause_count}, 合规: {st.session_state.compare_files[filename]['total_compliant']}/{st.session_state.compare_files[filename]['total_matched']})")
                else:
                    st.markdown(f"- {filename} (条款数: {clause_count})")
            with col_b:
                if st.button("分析", key=f"analyze_{filename}") and st.session_state.target_clauses:
                    matched_results, summary, total_compliant, total_matched = analyze_clause_matches(
                        st.session_state.target_clauses,
                        st.session_state.compare_files[filename]["clauses"],
                        st.session_state.api_key,
                        model_option
                    )
                    if matched_results is not None:
                        st.session_state.compare_files[filename]["matched_results"] = matched_results
                        st.session_state.compare_files[filename]["summary"] = summary
                        st.session_state.compare_files[filename]["total_compliant"] = total_compliant
                        st.session_state.compare_files[filename]["total_matched"] = total_matched
                        st.session_state.current_file = filename
                        st.success(f"✅ {filename} 分析完成，找到 {total_compliant} 条合规条款")

with col2:
    st.subheader("分析结果")
    
    # 显示文件选择标签
    if st.session_state.compare_files:
        st.markdown("**选择文件查看结果：**")
        cols_per_row = 3
        files = list(st.session_state.compare_files.items())
        rows = (len(files) + cols_per_row - 1) // cols_per_row
        
        for row in range(rows):
            cols = st.columns(cols_per_row)
            for col_idx in range(cols_per_row):
                file_idx = row * cols_per_row + col_idx
                if file_idx < len(files):
                    filename, data = files[file_idx]
                    with cols[col_idx]:
                        if "total_compliant" in data and data["total_compliant"] > 0:
                            status = f" ({data['total_compliant']}条合规)"
                        else:
                            status = ""
                        
                        if st.button(f"{filename.split('.')[0]}{status}", key=f"tab_{filename}"):
                            st.session_state.current_file = filename
    
    # 显示当前选中文件的分析结果
    if st.session_state.current_file:
        filename = st.session_state.current_file
        if filename in st.session_state.compare_files:
            file_data = st.session_state.compare_files[filename]
            matched_results = file_data.get("matched_results", None)
            summary = file_data.get("summary", "")
            total_compliant = file_data.get("total_compliant", 0)
            total_matched = file_data.get("total_matched", 0)
            
            if matched_results is not None:
                # 显示总体总结
                st.markdown("### 📊 总体分析总结")
                st.markdown('<div class="summary-box">', unsafe_allow_html=True)
                st.markdown(f"**匹配与合规概览：** 总匹配条款 {total_matched} 条，其中合规条款 {total_compliant} 条  \n")
                for para in re.split(r'\n+', summary):
                    if para.strip():
                        st.markdown(f"{para.strip()}  \n")
                st.markdown('</div>', unsafe_allow_html=True)
                
                # 显示合规条款的详细分析
                if matched_results:
                    st.markdown(f"### 🔍 合规条款详情 ({len(matched_results)} 条)")
                    
                    for clause_num, details in matched_results.items():
                        st.markdown(f'#### 目标条款第{details["target_num"]}条 vs 待比对条款第{details["compare_num"]}条')
                        st.markdown('<div class="matched-clause">', unsafe_allow_html=True)
                        
                        st.markdown("**目标条款内容：**")
                        st.write(details["target"][:500] + "..." if len(details["target"]) > 500 else details["target"])
                        
                        st.markdown("**待比对条款内容：**")
                        st.write(details["compare"][:500] + "..." if len(details["compare"]) > 500 else details["compare"])
                        
                        st.markdown('<div class="difference-section">', unsafe_allow_html=True)
                        st.markdown("**分析结果：**")
                        for para in re.split(r'\n+', details["analysis"]):
                            if para.strip():
                                st.markdown(f"{para.strip()}  \n")
                        st.markdown('</div>', unsafe_allow_html=True)
                        
                        st.markdown('</div>', unsafe_allow_html=True)
                
                # 生成并下载Word文档
                if target_file and matched_results is not None:
                    word_file = generate_word_document(
                        matched_results,
                        summary,
                        target_file.name,
                        filename,
                        total_compliant,
                        total_matched
                    )
                    
                    if word_file:
                        with open(word_file, "rb") as f:
                            st.download_button(
                                label=f"💾 下载 {filename} 的分析报告",
                                data=f,
                                file_name=f"政策条款比对报告_{filename}_{time.strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(word_file)
            else:
                st.info("请点击文件旁的'分析'按钮生成分析结果")
        else:
            st.warning("所选文件不存在，请重新选择")
    else:
        st.info("请上传待比对文件并选择一个文件查看分析结果")

# 帮助信息
with st.expander("ℹ️ 使用帮助"):
    st.markdown("""
    ### 工具特点
    1. **严格格式分割**：仅按照'一、二、三……'和'（一）（二）（三）……'两种格式分割条款
    2. **附件自动跳过**：自动识别并跳过PDF中的附件内容
    3. **双重匹配机制**：先按条款编号匹配，再按内容相似度匹配
    4. **合规性分析**：通过API判断条款合规性并生成分析报告
    
    ### 条款分割说明
    系统严格按照以下两种格式分割文件内容：
    - 格式1：中文数字+顿号，如"一、"、"二、"、"三、"等
    - 格式2：括号+中文数字+括号，如"（一）"、"（二）"、"（三）"等
    
    条款内容从当前格式标记开始，到下一个格式标记结束。
    
    ### 使用建议
    - 确保文档中的条款严格采用上述两种格式标记
    - 解析精度选择"严格"可过滤掉过短的条款内容
    - 分析完成后可下载Word格式报告保存分析结果
    """)
    
